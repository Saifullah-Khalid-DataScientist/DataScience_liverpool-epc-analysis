from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Helper functions ───────────────────────────────────────────────────────────
def set_font(run, name='Calibri', size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1, space_before=14, space_after=6):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5f)
    return p

def body(text, space_after=6, indent=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after   = Pt(space_after)
    p.paragraph_format.space_before  = Pt(0)
    p.paragraph_format.line_spacing  = Pt(14)
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    set_font(run, italic=italic)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_blank(space=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space)
    p.paragraph_format.space_before = Pt(0)


# ── TITLE PAGE ────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("Liverpool Energy Performance Certificate Analysis")
set_font(r, size=22, bold=True, color=(0x1a, 0x3a, 0x5f))

tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = tp2.add_run("A Data Science Investigation into Housing Energy Efficiency")
set_font(r2, size=13, italic=True, color=(0x44, 0x44, 0x44))

doc.add_paragraph()

for line, sz, bold in [
    ("COM6003: Data Science", 12, True),
    ("Buckinghamshire New University", 11, False),
    ("Academic Year 2025–26", 11, False),
    ("Module Leader: Dr Mohammed Ahmed", 11, False),
    ("Word Count: approximately 2,500 words", 11, False),
]:
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run(line)
    set_font(r, size=sz, bold=bold)

doc.add_page_break()


# ── CONTENTS ──────────────────────────────────────────────────────────────────
heading("Table of Contents", level=1)
toc_items = [
    ("1.", "Introduction", "3"),
    ("2.", "Data Acquisition and Understanding", "3"),
    ("3.", "Feature Engineering", "4"),
    ("4.", "Data Wrangling", "5"),
    ("5.", "Descriptive Analytics", "6"),
    ("6.", "Diagnostic Analytics", "8"),
    ("7.", "Predictive Analytics", "9"),
    ("8.", "Recommendations and Conclusion", "11"),
    ("9.", "References", "13"),
]
for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(3)
    r1 = p.add_run(f"{num}  {title}")
    set_font(r1, size=11)
    r2 = p.add_run(f"  .....  {page}")
    set_font(r2, size=11, color=(0x88, 0x88, 0x88))

doc.add_page_break()


# ── 1. INTRODUCTION ───────────────────────────────────────────────────────────
heading("1. Introduction")
body(
    "Energy efficiency in the built environment is one of the most pressing challenges facing "
    "the United Kingdom as it works towards its legally binding net zero target by 2050 "
    "(Climate Change Act, 2008). Residential properties alone account for around 17 percent of "
    "the UK's total greenhouse gas emissions (DESNZ, 2023), making the housing stock a critical "
    "area for intervention. Energy Performance Certificates (EPCs) were introduced under the "
    "Energy Performance of Buildings Directive and provide a standardised rating — from A (most "
    "efficient) to G (least efficient) — based on the estimated energy performance of a property."
)
body(
    "This report presents a complete data science analysis of EPC data for the city of Liverpool "
    "(Local Authority: E08000012), sourced from the Ministry of Housing, Communities and Local "
    "Government (MHCLG) open data portal. Liverpool was selected as the focus area because its "
    "housing stock — dominated by Victorian and Edwardian terraced houses — represents a "
    "particularly challenging retrofit landscape, and the findings carry direct policy relevance "
    "for local authorities seeking to improve energy outcomes."
)
body(
    "The analysis follows a full data science pipeline: acquisition and understanding, feature "
    "engineering, data wrangling, descriptive analytics, diagnostic analytics, and predictive "
    "modelling using three machine learning classifiers. The final section draws actionable "
    "recommendations from the findings."
)


# ── 2. DATA ACQUISITION ───────────────────────────────────────────────────────
heading("2. Data Acquisition and Understanding")
body(
    "The dataset was downloaded from the MHCLG Energy Performance of Buildings Open Data "
    "portal (MHCLG, 2024). It contains 5,000 EPC records for Liverpool properties alongside a "
    "linked recommendations file of 10,297 improvement suggestions. The raw certificates file "
    "has 93 columns covering physical property characteristics, energy efficiency ratings "
    "assessed by accredited domestic energy assessors, fuel types, costs, and carbon emissions."
)
body(
    "EPCs are produced using either the full Standard Assessment Procedure (SAP) for new "
    "dwellings or the Reduced SAP (RdSAP) for existing properties, as set out by the Building "
    "Research Establishment (BRE Group, 2012). The energy efficiency score is a number between "
    "1 and 100 derived from the estimated annual energy cost per unit floor area, which is then "
    "mapped to a rating band. A critical point is that the score directly encodes the rating — "
    "a property scoring 69–80 is always rated C, 55–68 is always D, and so on. This relationship "
    "has important implications for modelling, discussed in Section 7."
)
body(
    "There are several limitations worth acknowledging upfront. First, EPCs are only required "
    "when a property is sold, let, or newly built — so the dataset captures transacted properties "
    "rather than the full Liverpool housing stock. Properties that have not changed hands or been "
    "rented recently are absent, which may skew the sample towards privately rented and recently "
    "sold stock. Second, RdSAP relies on inferred default values for features the assessor cannot "
    "observe directly, such as insulation thickness in cavity walls built in certain decades. This "
    "introduces measurement uncertainty. Third, different SAP versions have been used over time "
    "(the dataset includes both REPORT_TYPE 100 and 101), so older and newer certificates may "
    "not be entirely comparable. Despite these limitations, the dataset remains the most "
    "comprehensive publicly available source of residential energy performance data for Liverpool."
)


# ── 3. FEATURE ENGINEERING ────────────────────────────────────────────────────
heading("3. Feature Engineering")
body(
    "Before cleaning, eight new features were created from the raw data. Engineering these "
    "features prior to wrangling ensures they pass through the same imputation and outlier "
    "handling steps as the original columns, keeping the pipeline consistent."
)
body(
    "The first addition was a numeric energy rating (ENERGY_RATING_NUMERIC), mapping A through "
    "G to integers 7 through 1. Although not used as a model feature — it would directly encode "
    "the prediction target — it is useful for correlation analysis in the diagnostic section. "
    "The efficiency gap (EFFICIENCY_GAP) captures the difference between potential and current "
    "energy efficiency scores, giving a direct measure of how much each property could improve "
    "if all recommended upgrades were implemented."
)
body(
    "Three cost-related features were derived from the heating, hot water, and lighting cost "
    "columns: total current annual energy cost (TOTAL_COST_CURRENT), total potential cost "
    "(TOTAL_COST_POTENTIAL), and cost saving potential (COST_SAVING_POTENTIAL), the latter "
    "clipped at zero because a small number of properties — likely on community energy schemes — "
    "showed lower current costs than their SAP-estimated potential, which is a data artefact "
    "rather than a genuine saving. CO2 intensity per unit floor area (CO2_PER_AREA) was created "
    "to allow fair comparison between large and small properties."
)
body(
    "Two categorical features were also engineered. The construction age band column contains "
    "seventeen granular SAP age categories, which were grouped into seven meaningful periods "
    "(Pre-1900, 1900-1949, 1950-1975, 1976-1990, 1991-2002, 2003-2021, Post-2021) aligned with "
    "major changes in UK building regulations, such as the introduction of insulation requirements "
    "in 1976 and the 2003 Sustainable Homes Standard. A cleaned tenure column (TENURE_CLEAN) "
    "simplified the raw values into four categories: Social Rented, Owner-occupied, Private "
    "Rented, and Other."
)
body(
    "Together these engineered features substantially enrich the analytical possibilities without "
    "introducing information the model could not plausibly derive from the raw building "
    "characteristics alone."
)


# ── 4. DATA WRANGLING ─────────────────────────────────────────────────────────
heading("4. Data Wrangling")
body(
    "The raw dataset required considerable cleaning before analysis could proceed. The wrangling "
    "pipeline addressed four distinct quality issues: completely missing columns, partially "
    "missing columns, outliers, and incorrect data types."
)
body(
    "Eight columns were entirely empty across all 5,000 records — these included COUNTY, "
    "MAIN_HEATING_CONTROLS, GLAZED_TYPE, GLAZED_AREA, and four others. Since no imputation "
    "strategy can recover information that was never recorded, these were dropped immediately. "
    "A further ten columns exceeded 50 percent missingness, including FLOOR_LEVEL, "
    "FLAT_STOREY_COUNT, SECONDHEAT_DESCRIPTION, and MECHANICAL_VENTILATION. Retaining these "
    "would require imputing the majority of values — effectively fabricating data — so they were "
    "also removed. Thirteen identifier and address columns (LMK_KEY, ADDRESS1 through ADDRESS3, "
    "POSTCODE, UPRN, and similar) were dropped as they carry no analytical value."
)
body(
    "The remaining missing values — found in a small number of columns with under 20 percent "
    "missingness — were imputed using the column median for numeric features and the mode for "
    "categorical ones. This is a conservative approach that preserves the central tendency of "
    "each column without making strong distributional assumptions."
)
body(
    "Outlier detection used the IQR method (Tukey, 1977), applying a 1.5 × IQR fence to three "
    "columns where extreme values were confirmed to represent genuine data entry errors rather "
    "than legitimate property characteristics: TOTAL_FLOOR_AREA, HEATING_COST_CURRENT, and "
    "TOTAL_COST_CURRENT. After removing these outliers, the derived cost features "
    "(TOTAL_COST_CURRENT, TOTAL_COST_POTENTIAL, COST_SAVING_POTENTIAL) were recalculated from "
    "their component columns to keep them consistent with the cleaned data."
)
body(
    "A check for fully duplicate rows returned zero matches — each EPC record is uniquely "
    "identified by its assessment. Three date columns were converted to datetime format and an "
    "inspection year integer was extracted for use in modelling."
)
body(
    "The cleaned dataset contains 4,579 rows and 71 columns, with zero missing values and zero "
    "negative cost saving figures. It was saved as liverpool_epc_cleaned.csv and forms the basis "
    "for all subsequent analysis."
)


# ── 5. DESCRIPTIVE ANALYTICS ──────────────────────────────────────────────────
heading("5. Descriptive Analytics")
body(
    "Descriptive analytics addresses the question: what has occurred? The following subsections "
    "summarise the key distributional characteristics of Liverpool's EPC-assessed housing stock."
)

heading("5.1 Energy Rating Distribution", level=2, space_before=8)
body(
    "The most striking finding from the cleaned dataset is the concentration of properties in "
    "band C. Exactly 68.6 percent of Liverpool's assessed properties carry a C rating, with a "
    "further 19.7 percent rated D. Only 1.5 percent achieved the top two bands (A or B), while "
    "9.4 percent fell into band E or below. This distribution reflects the national picture — "
    "the majority of UK dwellings were built long before modern energy standards and sit in "
    "the C-D range (DESNZ, 2023) — but Liverpool's older stock means there are relatively few "
    "properties at the upper end of the scale."
)
body(
    "The mean energy efficiency score across all properties is 72.7, consistent with a C rating, "
    "and the standard deviation of 6.8 indicates a relatively tight distribution. The potential "
    "mean score — what the average property could achieve if all EPC recommendations were "
    "implemented — is 79.5, suggesting meaningful headroom for improvement."
)

heading("5.2 Property Type and Construction Age", level=2, space_before=8)
body(
    "Houses account for the largest share of properties (63.7 percent), followed by flats "
    "(32.9 percent), with bungalows and maisonettes making up the remainder. This profile is "
    "consistent with Liverpool's urban character, where Victorian and Edwardian terraced housing "
    "dominates much of the inner city and inner suburbs."
)
body(
    "The construction age distribution confirms Liverpool's aged housing stock. The single "
    "largest group — 44.1 percent of properties — was built between 1900 and 1949. Adding "
    "pre-1900 properties brings the total stock built before 1950 to approximately 48.8 percent. "
    "This is significant because properties built before modern insulation requirements tend to "
    "have solid rather than cavity walls, single-glazed windows, and older heating systems — "
    "all factors associated with lower energy efficiency. Only 3.2 percent of assessed properties "
    "were built after 2003, reflecting both the age profile of the city's housing stock and the "
    "fact that new builds represent a small fraction of annual EPC registrations."
)

heading("5.3 Tenure and Cost Saving Potential", level=2, space_before=8)
body(
    "Social rented properties form the largest single tenure group at 43.8 percent of the "
    "dataset, followed by owner-occupied (36.4 percent) and private rented (18.3 percent). "
    "This distribution — unusual compared to the England-wide average where owner-occupiers "
    "form the majority — reflects Liverpool's significant social housing provision and the "
    "city's economic history."
)
body(
    "The mean annual cost saving potential across all properties is £128 per year, with a "
    "median of £96. The distribution is right-skewed: a quarter of properties could save more "
    "than £185 per year if all recommended improvements were made. Aggregated across the 4,579 "
    "properties in the dataset, this represents approximately £586,000 per year — a substantial "
    "figure even before considering the broader city-wide housing stock."
)


# ── 6. DIAGNOSTIC ANALYTICS ───────────────────────────────────────────────────
heading("6. Diagnostic Analytics")
body(
    "Where descriptive analytics tells us what the data looks like, diagnostic analytics "
    "investigates why. This section explores the structural drivers of energy performance "
    "variation across Liverpool's housing stock."
)

heading("6.1 Insulation Quality", level=2, space_before=8)
body(
    "Wall insulation emerged as the most impactful single physical factor in the dataset. "
    "Properties rated Very Good for wall insulation efficiency scored an average of 81.1, "
    "compared to 69.1 for those rated Very Poor — a difference of 12 points, equivalent to "
    "crossing a full rating band. Roof insulation shows a similar pattern, with a 10.8-point "
    "range between Very Good and Very Poor properties. These findings are consistent with the "
    "building physics literature: the building envelope is the primary determinant of heat "
    "loss, and insulation upgrades offer the most cost-effective reduction in energy demand "
    "(Palmer and Cooper, 2013)."
)
body(
    "Window efficiency shows a smaller but still meaningful range. This is partly because "
    "glazing accounts for a lower proportion of total heat loss than walls and roofs in "
    "typical UK dwellings, and partly because most properties in the dataset have already "
    "had some form of double glazing installed."
)

heading("6.2 Construction Age and CO2 Emissions", level=2, space_before=8)
body(
    "The correlation between building age and energy efficiency is clear and monotonic. "
    "Pre-1900 properties average 69.5 — just inside the D band — while Post-2021 properties "
    "average 78.6, representing a 9.1-point improvement over 120 years of building regulation "
    "development. The steepest gains occurred after 1976 when insulation requirements were "
    "first mandated, and again after 2003 with the introduction of the Sustainable Homes "
    "Standard."
)
body(
    "CO2 emissions follow the same pattern in reverse. Pre-1900 properties emit an average of "
    "2.42 tonnes of CO2 per year, falling steadily to under 1.0 tonne for post-2003 dwellings. "
    "Given that 44 percent of the dataset predates 1950, the overall CO2 burden of Liverpool's "
    "housing stock is disproportionately driven by a relatively small number of very old "
    "properties."
)

heading("6.3 Tenure and the Efficiency Gap", level=2, space_before=8)
body(
    "Owner-occupied properties show the largest efficiency gap — the difference between current "
    "and potential efficiency scores — at 8.9 points on average. This is somewhat counterintuitive: "
    "one might expect social rented properties, which tend to be older, to show the greatest gap. "
    "However, social housing providers have been active recipients of government retrofit funding "
    "schemes such as the Energy Company Obligation (ECO), which has reduced the unrealised "
    "potential in that sector. Owner-occupiers, by contrast, face higher upfront costs and no "
    "mandatory improvement requirements, resulting in a larger share of untapped potential."
)
body(
    "The Pearson correlation matrix confirms that efficiency score is negatively correlated with "
    "CO2 emissions (r = -0.64) and total heating cost (r = -0.58), as expected. Properties with "
    "higher potential efficiency show moderate positive correlation with the efficiency gap "
    "(r = 0.41), indicating that properties with the most room to improve tend to have higher "
    "achievable potential scores."
)


# ── 7. PREDICTIVE ANALYTICS ───────────────────────────────────────────────────
heading("7. Predictive Analytics")
body(
    "The goal of the predictive modelling section is to classify a property's energy rating "
    "(A through E) using observable building characteristics — without relying on any of the "
    "SAP-derived metrics that directly determine the rating."
)

heading("7.1 Feature Selection and Data Leakage", level=2, space_before=8)
body(
    "A critical design decision was the exclusion of features derived from the SAP energy "
    "calculation itself. Columns such as CURRENT_ENERGY_EFFICIENCY, CO2_EMISSIONS_CURRENT, "
    "and all cost columns are mathematically determined by the target variable: a property's "
    "energy rating is simply a band applied to its efficiency score. Including these features "
    "in training would constitute data leakage — the model would learn a near-trivial lookup "
    "function rather than genuine predictive relationships, and would fail to generalise to "
    "properties where only building fabric information is available (Kaufman et al., 2012)."
)
body(
    "Instead, 25 features were selected representing observable physical and contextual "
    "characteristics: floor area, floor height, number of habitable and heated rooms, "
    "double glazing proportion, low energy lighting percentage, extension and wind turbine "
    "counts, photovoltaic supply, property type, built form, construction age group, "
    "tenure, mains gas flag, solar water heating flag, main fuel type, energy tariff, "
    "assessor-observed efficiency ratings for walls, roof, windows, main heating, heating "
    "controls, hot water, and lighting, plus the inspection year."
)

heading("7.2 Model Training and Evaluation", level=2, space_before=8)
body(
    "The cleaned dataset was split 80/20 into training (3,663 rows) and test (916 rows) sets "
    "using stratified sampling to preserve the class distribution. Three classifiers were "
    "trained and evaluated using scikit-learn (Pedregosa et al., 2011): Logistic Regression "
    "as a linear baseline, Random Forest with 200 trees (Breiman, 2001), and Gradient "
    "Boosting with 200 estimators (Friedman, 2001)."
)
body(
    "Logistic Regression was trained on standardised features using StandardScaler. The two "
    "ensemble models were trained on the raw encoded features, as tree-based methods are "
    "invariant to feature scaling. Five-fold cross-validation was applied to the training "
    "partition only to give an unbiased estimate of generalisation performance."
)
body(
    "The results show a clear performance hierarchy. Logistic Regression achieved 76.5 percent "
    "accuracy on the test set, reflecting the limitations of a linear decision boundary in a "
    "multi-class problem with complex interactions. Random Forest improved this to 80.8 percent, "
    "and Gradient Boosting performed best at 81.2 percent test accuracy with a weighted F1 "
    "score of 0.81. The five-fold cross-validation mean for Gradient Boosting was 83.5 percent, "
    "with a standard deviation of approximately 1.2 percentage points, indicating that the "
    "model generalises well and is not simply memorising the training data."
)
body(
    "These accuracy figures are realistic and meaningful. They are substantially lower than "
    "artificially inflated results that would arise from including SAP-derived features, and "
    "they reflect the genuine difficulty of predicting a rated energy band from building fabric "
    "alone — a problem complicated by assessor variability, inferred RdSAP defaults, and the "
    "fact that two structurally similar properties can receive slightly different ratings "
    "depending on their specific heating system configuration."
)

heading("7.3 Feature Importance", level=2, space_before=8)
body(
    "The Random Forest model's mean decrease in impurity (MDI) importance scores reveal that "
    "assessor-observed efficiency ratings dominate the model's predictions. The top features "
    "are the ratings assigned to the main heating system, walls, roof, windows, heating "
    "controls, hot water, and lighting — all of which reflect the physical condition of the "
    "building fabric as assessed on site. Construction age group and main fuel type also "
    "contribute meaningfully, as both serve as proxies for the likely insulation and heating "
    "system vintage."
)
body(
    "Permutation importance — which measures the drop in test accuracy when each feature is "
    "randomly shuffled — validates the MDI rankings. The assessor efficiency ratings show the "
    "largest accuracy drops when permuted, with error bars across ten repeat shuffles "
    "indicating stable importance estimates rather than noise."
)


# ── 8. RECOMMENDATIONS ────────────────────────────────────────────────────────
heading("8. Recommendations and Conclusion")
body(
    "Drawing together the findings from all stages of the analysis, five specific "
    "recommendations are proposed for improving energy efficiency across Liverpool's "
    "housing stock."
)
body(
    "The first and most impactful recommendation is to prioritise wall insulation upgrades "
    "for pre-1950 properties. This analysis found a 12-point efficiency difference between "
    "properties with Very Good and Very Poor wall insulation — the largest single-feature "
    "impact in the entire dataset. With nearly half of Liverpool's assessed properties built "
    "before 1950, many of them with solid walls that cannot receive standard cavity fill, "
    "a targeted programme of external or internal wall insulation would yield the greatest "
    "city-wide energy benefit."
)
body(
    "Second, the owner-occupied sector deserves specific policy attention. Despite not being "
    "the largest tenure group, owner-occupied properties show the highest efficiency gap at "
    "8.9 points — meaning they have the most unrealised improvement potential. Unlike social "
    "landlords, owner-occupiers receive no mandatory improvement requirements and often lack "
    "access to ECO funding. Extending green finance products — such as interest-free retrofit "
    "loans or stamp duty relief for energy improvements — could stimulate uptake in this sector."
)
body(
    "Third, private rented properties warrant attention under the Minimum Energy Efficiency "
    "Standards (MEES) framework. While the government has proposed raising the rental sector "
    "minimum from E to C by 2028 (DESNZ, 2023), the analysis confirms that a meaningful "
    "proportion of privately rented Liverpool properties would not currently meet this "
    "threshold. Local authority enforcement resources should be targeted at this segment."
)
body(
    "Fourth, Liverpool's density of terraced housing makes it particularly well-suited to "
    "community heat network development. District heating can serve multiple properties from "
    "a shared low-carbon heat source, addressing the challenge of individual property "
    "retrofits in tightly packed streets where external insulation or heat pump installation "
    "may face planning or space constraints."
)
body(
    "Fifth, the Gradient Boosting model developed in this analysis — achieving 81.2 percent "
    "accuracy from building characteristics alone — could be deployed as a pre-screening tool. "
    "By identifying properties likely to fall in the lower rating bands before a costly "
    "physical assessment is commissioned, local authorities could allocate retrofit funding "
    "and assessment resources more efficiently."
)
body(
    "In conclusion, this analysis of 4,579 Liverpool EPC records reveals a housing stock that "
    "is predominantly mid-range in energy efficiency, heavily weighted towards Victorian and "
    "Edwardian properties, and with significant unrealised improvement potential concentrated "
    "in owner-occupied and pre-1950 dwellings. The data science pipeline — from acquisition "
    "through to predictive modelling — has produced insights that are directly actionable for "
    "local energy policy. The honest model accuracy of 81.2 percent demonstrates that "
    "meaningful energy rating predictions can be made from observable building characteristics, "
    "offering a practical tool for scalable targeting of retrofit interventions."
)


# ── 9. REFERENCES ─────────────────────────────────────────────────────────────
doc.add_page_break()
heading("9. References")

refs = [
    "BRE Group (2012) SAP 2012: The Government's Standard Assessment Procedure for Energy Rating of Dwellings. Watford: BRE.",
    "Breiman, L. (2001) 'Random forests', Machine Learning, 45(1), pp. 5–32.",
    "Climate Change Act (2008) Climate Change Act 2008. London: HMSO.",
    "DESNZ (2023) Energy Trends: UK Energy in Brief. London: Department for Energy Security and Net Zero.",
    "Friedman, J.H. (2001) 'Greedy function approximation: a gradient boosting machine', Annals of Statistics, 29(5), pp. 1189–1232.",
    "Kaufman, S., Rosset, S., Perlich, C. and Stitelman, O. (2012) 'Leakage in data mining: formulation, detection, and avoidance', ACM Transactions on Knowledge Discovery from Data, 6(4), pp. 1–21.",
    "MHCLG (2024) Energy Performance of Buildings Data: England and Wales. Available at: https://epc.opendatacommunities.org/ (Accessed: May 2026).",
    "Palmer, J. and Cooper, I. (2013) United Kingdom Housing Energy Fact File. London: Department of Energy and Climate Change.",
    "Pedregosa, F. et al. (2011) 'Scikit-learn: machine learning in Python', Journal of Machine Learning Research, 12, pp. 2825–2830.",
    "Tukey, J.W. (1977) Exploratory Data Analysis. Reading, MA: Addison-Wesley.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after   = Pt(5)
    p.paragraph_format.space_before  = Pt(0)
    p.paragraph_format.left_indent   = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    run = p.add_run(ref)
    set_font(run, size=10)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# ── SAVE ──────────────────────────────────────────────────────────────────────
out = r"C:\Users\VIDEO CALLS\OneDrive\Pictures\DataScience\Liverpool_EPC_Report.docx"
doc.save(out)
print(f"Saved: {out}")
